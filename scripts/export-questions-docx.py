"""Export active questions to DOCX (reads output/doc/questions-export.json)."""
from __future__ import annotations

import json
import os
import sys
from datetime import datetime
from pathlib import Path
from urllib.parse import urljoin

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
JSON_FILE = ROOT / "data" / "questions-canonical.json"
OUTPUT_DIR = ROOT / "output" / "doc"
OUTPUT_FILE = OUTPUT_DIR / "Split_Second_214_Sorular.docx"

SLOT_LABELS = {
    "morning": "Sabah (08:00)",
    "afternoon": "Öğle (14:00)",
    "evening": "Akşam (20:00)",
}


def fetch_questions_from_supabase() -> list[dict]:
    url = os.environ.get("EXPO_PUBLIC_SUPABASE_URL", "").rstrip("/")
    key = os.environ.get("EXPO_PUBLIC_SUPABASE_ANON_KEY", "")
    if not url or not key:
        raise RuntimeError("Set EXPO_PUBLIC_SUPABASE_URL and EXPO_PUBLIC_SUPABASE_ANON_KEY")

    endpoint = urljoin(url + "/", "rest/v1/questions")
    params = {
        "select": "scheduled_date,time_slot,category,question_text,question_text_tr,option_a,option_a_tr,option_b,option_b_tr",
        "is_active": "eq.true",
        "order": "scheduled_date.asc,time_slot.asc",
        "limit": "500",
    }
    import urllib.request

    query = "&".join(f"{k}={v}" for k, v in params.items())
    req = urllib.request.Request(
        f"{endpoint}?{query}",
        headers={
            "apikey": key,
            "Authorization": f"Bearer {key}",
        },
    )
    with urllib.request.urlopen(req, timeout=60) as resp:
        data = json.loads(resp.read().decode("utf-8"))

    slot_order = {"morning": 0, "afternoon": 1, "evening": 2}
    data.sort(
        key=lambda q: (q["scheduled_date"], slot_order.get(q["time_slot"], 9)),
    )
    return data


def load_questions(path: Path) -> list[dict]:
    if path.is_file():
        data = json.loads(path.read_text(encoding="utf-8"))
        if isinstance(data, dict) and "questions" in data:
            return data["questions"]
        return data
    return fetch_questions_from_supabase()


def add_para(doc: Document, label: str, value: str | None) -> None:
    p = doc.add_paragraph()
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    p.add_run(value if value else "(yok)")


def build_doc(questions: list[dict]) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    title = doc.add_heading("Split Second - Soru Envanteri", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        f"Toplam {len(questions)} aktif soru | "
        f"Olusturulma: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    ).font.size = Pt(10)

    doc.add_paragraph()

    for index, q in enumerate(questions, start=1):
        date = q.get("scheduled_date", "")
        slot = q.get("time_slot", "")
        slot_label = SLOT_LABELS.get(slot, slot)
        category = q.get("category", "")

        doc.add_heading(f"{index}. {date} - {slot_label}", level=2)
        add_para(doc, "Kategori", category)

        add_para(doc, "Soru (EN)", q.get("question_text"))
        add_para(doc, "Soru (TR)", q.get("question_text_tr"))
        add_para(doc, "A (EN)", q.get("option_a"))
        add_para(doc, "A (TR)", q.get("option_a_tr"))
        add_para(doc, "B (EN)", q.get("option_b"))
        add_para(doc, "B (TR)", q.get("option_b_tr"))

        if index < len(questions):
            doc.add_paragraph("_" * 48)

    return doc


def main() -> int:
    json_path = Path(sys.argv[1]) if len(sys.argv) > 1 else JSON_FILE

    # Load .env from project root for Supabase fetch
    env_path = ROOT / ".env"
    if env_path.is_file():
        for line in env_path.read_text(encoding="utf-8").splitlines():
            line = line.strip()
            if line and not line.startswith("#") and "=" in line:
                k, _, v = line.partition("=")
                os.environ.setdefault(k.strip(), v.strip())

    questions = load_questions(json_path)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    if not json_path.is_file():
        json_path.write_text(json.dumps(questions, ensure_ascii=False, indent=2), encoding="utf-8")

    out_name = f"Split_Second_{len(questions)}_Sorular.docx"
    out_path = OUTPUT_DIR / out_name
    doc = build_doc(questions)
    doc.save(out_path)
    print(f"Wrote {len(questions)} questions to {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
